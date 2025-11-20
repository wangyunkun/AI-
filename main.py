import flet as ft
import base64
import json
import threading
import os
import sys
import io
from datetime import datetime
from openai import OpenAI
# 【新增】导入 docx 库
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ================= 1. 预设配置 =================
PROVIDER_PRESETS = {
    "阿里百炼 (Alibaba)": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "model": "qwen-vl-max",
        "api_key": ""
    },
    "硅基流动 (SiliconFlow)": {
        "base_url": "https://api.siliconflow.cn/v1",
        "model": "Qwen/Qwen2-VL-72B-Instruct",
        "api_key": ""
    },
    "DeepSeek (官方)": {
        "base_url": "https://api.deepseek.com/v1",
        "model": "deepseek-chat",
        "api_key": ""
    },
    "自定义 (Custom)": {
        "base_url": "",
        "model": "",
        "api_key": ""
    }
}

DEFAULT_PROMPT = """你是一位拥有30年一线经验的**国家注册安全工程师**及**工程质量监理专家**。你的眼神如鹰隼般锐利，绝不放过任何一个细微的安全隐患或违规施工行为。

你的任务是审查施工现场照片，重点针对**施工机械使用**、**施工工艺规范**以及**通用EHS风险**进行全方位扫描。

请按照以下逻辑顺序，对画面进行“像素级”的排查：

### 第一优先级：大型机械与特种设备（深度审查）
1. **起重吊装**：
   - 汽车吊/履带吊：支腿是否完全伸出并垫实？吊臂下是否有人员逗留？是否有司索工/指挥人员？
   - 吊装作业设备：是否违章用装载机、挖机等机械进行吊装？是否有违规起吊（歪拉斜吊、超载、非标准吊具）？
2. **土方机械**：
   - 挖掘机/装载机：作业半径内是否有闲杂人员？驾驶室是否有人违规搭乘？停放位置是否在大坡度或坑边？
3. **桩机/钻机**：
   - 设备是否稳固？电缆是否拖地浸水？

### 第二优先级：施工工艺与临时设施（专业审查）
1. **脚手架与模板支撑**：
   - 立杆是否垂直？是否有扫地杆？剪刀撑是否连续设置？脚手板是否铺满且固定？
   - **违规判定**：严禁钢管与木方混用、严禁缺少底座。
2. **临电作业**：
   - 是否落实“一机一闸一漏一箱”？配电箱门是否关闭？电缆是否乱拉乱接或经过通道未防护？
3. **高处作业与临边**：
   - “四口五临边”是否有防护栏杆？安全网是否挂设严密？作业平台是否稳固？

### 第三优先级：人员行为与文明施工（基础审查）
1. **个人防护 (PPE)**：
   - 安全帽（必须系下颌带）、反光衣、高处作业必须系挂五点式安全带（高挂低用）。
2. **消防与动火**：
   - 气瓶是否防倾倒？氧气/乙炔间距是否足够（5米）？动火点旁是否有灭火器？是否有接火斗？
3. **文明施工**：
   - 材料是否分类堆放？裸土是否覆盖？路面是否积水或泥泞？

---

### 输出规则（极其重要）

1. **引用标准**：在指出问题时，请尽量匹配最精确的中国国标或行标。
   - 机械类参考：GB 6067《起重机械安全规程》、JGJ 33《建筑机械使用安全技术规程》。
   - 施工类参考：JGJ 59《建筑施工安全检查标准》、JGJ 130《扣件式钢管脚手架安全技术规范》、GB 50194《建设工程施工现场供用电安全规范》。
2. **数量统计**：如果同一类问题出现多次（如3人未戴头盔），请合并为一条，但要说明数量。
3. **宁严勿漏**：对于模糊不清的隐患，用“疑似”字样指出，提示人工复核。

请返回纯净的 JSON 列表（无 Markdown 标记），格式如下：
[
    {
        "issue": "【机械】挖掘机作业半径内有2名工人违规穿越，且无人指挥",
        "regulation": "违反《建筑机械使用安全技术规程》JGJ 33-2012 第x条",
        "correction": "立即停止作业，设置警戒隔离区，配备专职指挥人员"
    },
    {
        "issue": "【工艺】落地式脚手架纵向剪刀撑未连续设置，且立杆悬空",
        "regulation": "违反《建筑施工扣件式钢管脚手架安全技术规范》JGJ 130-2011",
        "correction": "立即整改，补齐剪刀撑，立杆底部增设垫板和底座"
    }
]

如果未发现任何问题，返回 []。
"""

CONFIG_FILE = "app_config_final.json"


class SafetyApp:
    def __init__(self):
        self.config = self.load_config()
        self.current_image_path = None
        self.current_data = []
        self.client = None

    def load_config(self):
        default = {"current_provider": "阿里百炼 (Alibaba)", "system_prompt": DEFAULT_PROMPT,
                   "providers": PROVIDER_PRESETS}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    saved = json.load(f)
                    if "providers" not in saved:
                        saved["providers"] = PROVIDER_PRESETS
                    else:
                        for k, v in PROVIDER_PRESETS.items():
                            if k not in saved["providers"]: saved["providers"][k] = v
                    return saved
            except:
                return default
        return default

    def save_config_to_file(self):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
        except:
            pass

    def init_client(self):
        p = self.config.get("current_provider")
        conf = self.config["providers"].get(p, {})
        if conf.get("api_key") and conf.get("base_url"):
            self.client = OpenAI(api_key=conf["api_key"], base_url=conf["base_url"])
            return True
        return False


def main(page: ft.Page):
    # ================= 页面设置 =================
    page.title = "普洱版纳区域安全检查AI"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.bgcolor = "#f2f4f7"
    page.padding = 0
    page.scroll = ft.ScrollMode.AUTO

    page.window.width = 1200
    page.window.height = 850
    page.window.min_width = 380
    page.window.min_height = 600

    app = SafetyApp()

    # ================= 详情抽屉 =================
    def show_bottom_sheet(item):
        bs_content.controls = [
            ft.Container(height=10),
            ft.Container(width=40, height=5, bgcolor=ft.Colors.GREY_300, border_radius=10,
                         alignment=ft.alignment.center),
            ft.Text("隐患详情", size=18, weight="bold", text_align="center"),
            ft.Divider(),
            ft.Text("⚠️ 隐患描述", color="red", weight="bold"),
            ft.Container(content=ft.Text(item.get("issue", ""), selectable=True), padding=10, bgcolor=ft.Colors.RED_50,
                         border_radius=6),
            ft.Container(height=10),
            ft.Text("⚖️ 依据规范", color="blue", weight="bold"),
            ft.Container(content=ft.Text(item.get("regulation", ""), selectable=True), padding=10,
                         bgcolor=ft.Colors.BLUE_50, border_radius=6),
            ft.Container(height=10),
            ft.Text("🛠️ 整改建议", color="green", weight="bold"),
            ft.Container(content=ft.Text(item.get("correction", ""), selectable=True), padding=10,
                         bgcolor=ft.Colors.GREEN_50, border_radius=6),
            ft.Container(height=30)
        ]
        bs.open = True
        page.update()

    bs_content = ft.Column(scroll=ft.ScrollMode.AUTO, tight=True)
    bs = ft.BottomSheet(content=ft.Container(content=bs_content, padding=20,
                                             border_radius=ft.border_radius.only(top_left=15, top_right=15)),
                        dismissible=True)
    page.overlay.append(bs)

    # ================= 列表渲染 =================
    result_column = ft.Column(spacing=10)

    def render_results(data):
        result_column.controls.clear()
        if not data:
            result_column.controls.append(
                ft.Container(content=ft.Text("暂无数据，请上传图片分析", color="grey"), alignment=ft.alignment.center,
                             padding=30))
        else:
            for i, item in enumerate(data):
                card = ft.Container(
                    bgcolor="white", padding=15, border_radius=10,
                    shadow=ft.BoxShadow(blur_radius=5, color=ft.Colors.BLACK12),
                    on_click=lambda e, d=item: show_bottom_sheet(d),
                    content=ft.Column([
                        ft.Row([
                            ft.Icon(ft.Icons.WARNING_ROUNDED, color="red"),
                            ft.Text(f"隐患 #{i + 1}", weight="bold", size=16),
                            ft.Container(expand=True),
                            ft.Icon(ft.Icons.ARROW_FORWARD_IOS, size=14, color="grey")
                        ]),
                        ft.Text(item.get("issue", ""), max_lines=2, overflow=ft.TextOverflow.ELLIPSIS),
                        ft.Divider(height=5, color="transparent"),
                        ft.Text(item.get("regulation", "")[:20] + "...", size=12, color="grey")
                    ])
                )
                result_column.controls.append(card)
        page.update()

    # ================= UI 组件 =================
    status_txt = ft.Text("请配置 Key", color="grey", size=12)
    img_control = ft.Image(src="https://placehold.co/600x400?text=Preview", fit=ft.ImageFit.CONTAIN, expand=True,
                           border_radius=8)
    img_container = ft.Container(content=img_control, height=250, bgcolor=ft.Colors.BLACK12, border_radius=10,
                                 alignment=ft.alignment.center)

    # ================= 核心逻辑 =================
    def save_config(e):
        p = dd_provider.value
        app.config["current_provider"] = p
        app.config["system_prompt"] = tf_prompt.value
        app.config["providers"][p]["base_url"] = tf_url.value.strip()
        app.config["providers"][p]["model"] = tf_model.value.strip()
        app.config["providers"][p]["api_key"] = tf_key.value.strip()
        app.save_config_to_file()
        status_txt.value = "✅ 配置已保存"
        page.close(dlg_settings)
        page.update()

    def refresh_settings(val):
        conf = app.config["providers"].get(val, {})
        tf_url.value = conf.get("base_url", "")
        tf_model.value = conf.get("model", "")
        tf_key.value = conf.get("api_key", "")
        page.update()

    def on_exit_app(e):
        try:
            page.window.close()
        except:
            sys.exit(0)

    def run_task(e):
        if not app.init_client():
            status_txt.value = "❌ 未配置API"
            status_txt.color = "red"
            page.update()
            return
        btn_analyze.disabled = True
        btn_analyze.text = "分析中..."
        page.update()

        def task():
            try:
                p = app.config["current_provider"]
                with open(app.current_image_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                resp = app.client.chat.completions.create(
                    model=app.config["providers"][p]["model"],
                    messages=[{"role": "system", "content": app.config["system_prompt"]},
                              {"role": "user",
                               "content": [{"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                                           {"type": "text", "text": "找出所有隐患"}]}], temperature=0.1
                )
                content = resp.choices[0].message.content.replace("```json", "").replace("```", "")
                s, e_idx = content.find('['), content.rfind(']') + 1
                data = json.loads(content[s:e_idx]) if s != -1 and e_idx != -1 else []
                app.current_data = data
                render_results(data)
                status_txt.value = "✅ 分析完成"
                status_txt.color = "green"
                btn_analyze.text = "重新分析"
                btn_analyze.disabled = False
                btn_export.disabled = False
                page.update()
            except Exception as err:
                status_txt.value = f"❌ {str(err)[:20]}"
                status_txt.color = "red"
                btn_analyze.disabled = False
                page.update()

        threading.Thread(target=task).start()

    def on_picked(e):
        if e.files:
            app.current_image_path = e.files[0].path
            img_control.src = e.files[0].path
            status_txt.value = "📸 图片已就绪"
            status_txt.color = "blue"
            btn_analyze.disabled = False
            page.update()

    # ================= 【核心】Word 导出逻辑 =================
    def on_save_word(e):
        """
        Word 导出函数
        """

        # --- 内部函数：生成 Word 二进制流 ---
        def generate_word_bytes():
            if not app.current_data:
                raise Exception("无数据可导出")

            # 1. 创建文档
            doc = Document()

            # 设置中文字体支持（可选，防止乱码）
            doc.styles['Normal'].font.name = u'宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            # 2. 添加大标题
            heading = doc.add_heading('AI 安全隐患排查报告', 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"隐患数量: {len(app.current_data)} 项")
            doc.add_paragraph("-" * 30)

            # 3. 创建表格 (行数 = 数据行数 + 1表头, 列数=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'  # 添加边框

            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '隐患描述'
            hdr_cells[1].text = '依据规范'
            hdr_cells[2].text = '整改建议'

            # 填充数据
            for item in app.current_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("issue", "")
                row_cells[1].text = item.get("regulation", "")
                row_cells[2].text = item.get("correction", "")

            # 4. 保存到内存流
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            return output_buffer.getvalue()

        # -------------------------------------------

        is_mobile = page.platform in ["android", "ios"]

        try:
            # 准备文件名 (.docx)
            filename = f"安全报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            # === 分支 A：电脑端 ===
            if not is_mobile:
                if hasattr(e, "path") and e.path:
                    save_path = e.path
                    if not save_path.endswith(".docx"): save_path += ".docx"

                    data = generate_word_bytes()
                    with open(save_path, "wb") as f:
                        f.write(data)

                    page.snack_bar = ft.SnackBar(ft.Text(f"✅ 导出Word成功"), bgcolor="green")
                    page.snack_bar.open = True
                    page.update()
                return

            # === 分支 B：手机端 ===
            word_data = generate_word_bytes()

            # 安卓路径尝试列表
            possible_paths = [
                f"/storage/emulated/0/Download/{filename}",
                f"/sdcard/Download/{filename}",
                os.path.join(os.getcwd(), filename)
            ]

            success_path = ""
            error_log = []

            for path in possible_paths:
                try:
                    with open(path, "wb") as f:
                        f.write(word_data)
                    success_path = path
                    break
                except Exception as e_path:
                    error_log.append(f"{path}: {str(e_path)}")
                    continue

            if success_path:
                dlg_success = ft.AlertDialog(
                    title=ft.Text("导出成功"),
                    content=ft.Text(f"Word报告已保存至手机【下载/Download】文件夹。\n\n文件名:\n{filename}", size=16),
                    actions=[ft.TextButton("确定", on_click=lambda e: page.close(dlg_success))]
                )
                page.open(dlg_success)
                page.update()
            else:
                # 如果都失败了，打印具体路径错误，方便排查
                all_errors = "\n".join(error_log)
                raise Exception(f"写入失败，请检查权限。\n{all_errors}")

        except Exception as err:
            page.snack_bar = ft.SnackBar(ft.Text(f"导出失败: {str(err)}"), bgcolor="red")
            page.snack_bar.open = True
            page.update()

    # ================= 布局组装 =================
    dd_provider = ft.Dropdown(label="厂商", options=[ft.dropdown.Option(k) for k in PROVIDER_PRESETS],
                              value=app.config.get("current_provider"),
                              on_change=lambda e: refresh_settings(e.control.value))
    tf_key = ft.TextField(label="Key", password=True)
    tf_url = ft.TextField(label="URL")
    tf_model = ft.TextField(label="Model")
    tf_prompt = ft.TextField(label="提示词", value=app.config.get("system_prompt"), multiline=True, min_lines=3)
    refresh_settings(app.config.get("current_provider"))

    dlg_settings = ft.AlertDialog(title=ft.Text("设置"),
                                  content=ft.Column([dd_provider, tf_key, tf_url, tf_model, tf_prompt],
                                                    scroll=ft.ScrollMode.AUTO, height=350, width=300),
                                  actions=[ft.TextButton("保存", on_click=save_config)])

    pick_dlg = ft.FilePicker(on_result=on_picked)
    save_dlg = ft.FilePicker(on_result=on_save_word)  # 更改为 Word 保存

    page.overlay.extend([pick_dlg, save_dlg])

    header = ft.Container(
        content=ft.Row([
            ft.Text("🛡️ 普洱版纳区域安全检查AI", size=18, weight="bold"),
            ft.Row([
                ft.IconButton(ft.Icons.SETTINGS, tooltip="设置", on_click=lambda e: page.open(dlg_settings)),
                ft.IconButton(ft.Icons.EXIT_TO_APP, tooltip="退出系统", icon_color="red", on_click=on_exit_app)
            ])
        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
        padding=15, bgcolor="white", border_radius=10, shadow=ft.BoxShadow(blur_radius=2, color=ft.Colors.BLACK12)
    )

    btn_style = ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=8), padding=15)
    btn_upload = ft.ElevatedButton("选图", icon=ft.Icons.IMAGE, on_click=lambda _: pick_dlg.pick_files(),
                                   style=btn_style)
    btn_analyze = ft.ElevatedButton("分析", icon=ft.Icons.AUTO_AWESOME, on_click=run_task, disabled=True,
                                    style=ft.ButtonStyle(bgcolor="blue", color="white", padding=15,
                                                         shape=ft.RoundedRectangleBorder(radius=8)))

    default_filename = f"安全报告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    def trigger_export(e):
        if page.platform in ["android", "ios"]:
            on_save_word(None)
        else:
            # 电脑端限制为 docx
            save_dlg.save_file(file_name=default_filename, allowed_extensions=["docx"])

    btn_export = ft.ElevatedButton("导出报告", icon=ft.Icons.DESCRIPTION,  # 图标改为文档
                                   on_click=trigger_export, disabled=True,
                                   style=ft.ButtonStyle(color="purple", padding=15,  # 颜色改为紫色
                                                        shape=ft.RoundedRectangleBorder(radius=8)))

    layout = ft.ResponsiveRow([
        ft.Column(col={"xs": 12, "md": 5}, controls=[
            ft.Container(content=img_container, bgcolor="white", padding=10, border_radius=10),
            ft.Container(height=5),
            ft.Row([
                ft.Column([btn_upload], expand=1),
                ft.Column([btn_analyze], expand=1),
                ft.Column([btn_export], expand=1),
            ]),
            ft.Container(content=status_txt, alignment=ft.alignment.center),
        ]),

        ft.Column(col={"xs": 12, "md": 7}, controls=[
            ft.Container(
                content=ft.Column([
                    ft.Text("📋 检查结果", size=16, weight="bold", color=ft.Colors.GREY_700),
                    result_column
                ]),
                bgcolor="white", padding=15, border_radius=10
            )
        ])
    ], spacing=20)

    page.add(
        ft.SafeArea(
            ft.Container(
                content=ft.Column([
                    header,
                    layout
                ]),
                padding=10
            )
        )
    )

    render_results([])


ft.app(target=main)
